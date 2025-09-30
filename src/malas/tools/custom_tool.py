from crewai.tools import BaseTool
from typing import Type
from pydantic import BaseModel, Field
from langchain_community.tools import DuckDuckGoSearchResults
import requests
import trafilatura

from docx import Document
import tempfile

from malas.crews.models.TaskOutput import Outline, ReferenceItem, References, Subbab


class DuckDuckGoToolInput(BaseModel):
    """Input schema for DuckDuckGoSearchTool."""
    query: str = Field(..., description="The search query to run on DuckDuckGo.")


class DuckDuckGoSearchTool(BaseTool):
    name: str = "DuckDuckGo Search Tool"
    description: str = (
        "Useful for searching information on the internet using DuckDuckGo. "
        "Provide a query and it will return top search results."
    )
    args_schema: Type[BaseModel] = DuckDuckGoToolInput

    def _run(self, query: str) -> list[dict]:
        search = DuckDuckGoSearchResults(max_results=3)
        raw = search.run(query)  # string mentah
        results = []
    
        # contoh sederhana split tiap line, parse title, url, snippet
        for line in raw.split("\n"):
            if "title:" in line and "link:" in line:
                results.append({
                    "judul": line.split("title:")[1].split(",")[0].strip(),
                    "url": line.split("link:")[1].strip(),
                    "tahun": None,  # bisa parse dari snippet jika ada
                    "penulis": None
                })
        return results


class ReferenceFinderTool(BaseTool):
    name: str = "Reference Finder Tool"
    description: str = (
        "Gunakan untuk mencari referensi akademik (artikel jurnal, prosiding, buku) "
        "berdasarkan topik atau kata kunci."
    )

    def _run(self, query: str) -> str:
        url = "https://api.crossref.org/works"
        params = {"query": query, "rows": 5}  # ambil 5 teratas
        try:
            res = requests.get(url, params=params, timeout=10)
            res.raise_for_status()
            items = res.json()["message"]["items"]

            results = []
            for item in items:
                title = item.get("title", ["No title"])[0]
                author = ", ".join([a.get("family", "") for a in item.get("author", [])[:3]])
                year = item.get("issued", {}).get("date-parts", [[None]])[0][0]
                doi = item.get("DOI", "")
                results.append(f"{title} ({year}) - {author} | DOI: {doi}")

            return "\n".join(results) if results else "Tidak ada referensi ditemukan."
        except Exception as e:
            return f"Error mencari referensi: {e}"
        


class ResearchExtractorTool(BaseTool):
    name: str = "Research Extractor Tool"
    description: str = (
        "Ekstrak konten dari artikel online (misalnya berita, jurnal, blog). "
        "Masukkan URL artikel."
    )

    def _run(self, url: str) -> str:
        try:
            downloaded = trafilatura.fetch_url(url)
            text = trafilatura.extract(downloaded)
            return text[:500] if text else "Konten tidak bisa diekstrak."
        except Exception as e:
            return f"Error mengambil konten: {e}"
        



class ExportDocxOutput(BaseModel):
    file_path: str  # simpan path file docx

class ExportDocxTool(BaseTool):
    name: str = "export_docx_tool"
    description: str = (
        "Ekstrak konten dari makalah (dict) menjadi file docx."
    )

    def _run(self, makalah: dict) -> ExportDocxOutput:
        """
        Terima makalah (output reviewer dalam bentuk dict),
        buat file .docx dari semua bab dan referensi.
        """
        doc = Document()
        doc.add_heading(makalah.get("judul", ""), level=0)

        for bab in makalah.get("bab", []):
            doc.add_heading(bab.get("judul", ""), level=1)
            if "konten" in bab and bab["konten"]:
                doc.add_paragraph(bab["konten"])
            
            # SubBab
            for sub in bab.get("subbab", []):
                doc.add_heading(sub.get("judul", ""), level=2)
                if "konten" in sub and sub["konten"]:
                    doc.add_paragraph(sub["konten"])
                if "referensi" in sub and sub["referensi"]:
                    doc.add_paragraph("Referensi:")
                    for r in sub["referensi"]:
                        doc.add_paragraph(f"- {r}", style="List Bullet")
            
            # Referensi Bab
            if "referensi" in bab and bab["referensi"]:
                doc.add_paragraph("Referensi:")
                for r in bab["referensi"]:
                    doc.add_paragraph(f"- {r}", style="List Bullet")

        # Simpan ke file sementara
        tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp_file.name)

        return ExportDocxOutput(file_path=tmp_file.name)



class MockOutlineTool(BaseTool):
    name: str = "Mock Outline Generator"
    description: str = "Generates a static, predefined makalah outline for debugging."

    def _run(self) -> Outline:
        print("\n--- RUNNING MOCK FOR OUTLINE TASK ---\n")
        # Mengembalikan objek Pydantic, CrewAI akan menanganinya
        return Outline(
            subbabs={
                "BAB I: PENDAHULUAN": Subbab(sections=["Latar Belakang", "Rumusan Masalah", "Tujuan Penelitian"]),
                "BAB II: TINJAUAN PUSTAKA": Subbab(sections=["Penelitian Terdahulu", "Landasan Teori"]),
                "BAB III: METODOLOGI PENELITIAN": Subbab(sections=["Jenis Penelitian", "Sumber Data", "Teknik Analisis Data"])
            }
        )
    


class MockReferenceTool(BaseTool):
    name: str = "Mock Reference Searcher"
    description: str = "Generates a static, predefined list of references for debugging."

    def _run(self) -> References:
        print("\n--- RUNNING MOCK FOR REFERENCE TASK ---\n")
        # Mengembalikan objek Pydantic
        return References(
            references=[
                ReferenceItem(title="Artificial Intelligence: A Modern Approach", authors=["Stuart Russell", "Peter Norvig"], year=2020, link="http://aima.cs.berkeley.edu/"),
                ReferenceItem(title="Deep Learning", authors=["Ian Goodfellow", "Yoshua Bengio", "Aaron Courville"], year=2016)
            ]
        )